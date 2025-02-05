import os
import io
import re
import json
import pytesseract
import pdfplumber
import fitz  # PyMuPDF
import pandas as pd
from PIL import Image
from docx import Document

from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build

from ultralytics import YOLO
import openai  # (Optional) If you want AI-based refinement for YOLO

# -------------------------------------------------------------------
# 1) CONFIGURATION
# -------------------------------------------------------------------

SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
EXTRACTION_OUTPUT_FOLDER = "extracted_output"
DOWNLOAD_FOLDER = "downloads"
PROCESSED_FILES_LOG = "processed_files.json"

# Tesseract multi-language (adjust as needed)
TESSERACT_LANGS = "eng"

# YOLO (if needed for object detection)
model = YOLO("yolov8n.pt")

# Optional OpenAI config for refining YOLO output
openai.api_key = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = "gpt-3.5-turbo"

# -------------------------------------------------------------------
# 2) AUTHENTICATION / DRIVE FETCH
# -------------------------------------------------------------------

def authenticate_google_drive():
    """
    OAuth2 authentication with Google for Drive read-only access.
    """
    flow = InstalledAppFlow.from_client_secrets_file("credentials.json", SCOPES)
    creds = flow.run_local_server(port=8080)
    return creds

def fetch_drive_files(creds):
    """
    Paginated retrieval of all files from Google Drive.
    Also request 'webViewLink' for a direct link to the file (source_url).
    """
    service = build("drive", "v3", credentials=creds)
    all_files = []
    page_token = None

    while True:
        response = service.files().list(
            pageSize=100,
            pageToken=page_token,
            fields="nextPageToken, files(id, name, mimeType, modifiedTime, webViewLink)"
        ).execute()

        files = response.get("files", [])
        all_files.extend(files)

        page_token = response.get("nextPageToken")
        if not page_token:
            break

    return all_files, service

def download_file(service, file_id, file_name):
    """
    Download the file from Drive to the local DOWNLOAD_FOLDER.
    """
    request = service.files().get_media(fileId=file_id)
    file_path = os.path.join(DOWNLOAD_FOLDER, file_name)
    with open(file_path, "wb") as f:
        f.write(request.execute())
    return file_path

# -------------------------------------------------------------------
# 3) EXTRACTION UTILITIES
# -------------------------------------------------------------------

def regex_link_extractor(text):
    """Fallback regex to find URLs in any raw text."""
    pattern = r'(https?://[^\s]+)'
    found = re.findall(pattern, text)
    return list(set(found))  # unique

def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()

def extract_images_from_pdf(pdf_path):
    """
    Extract images from PDF using PyMuPDF, then OCR each image.
    """
    doc = fitz.open(pdf_path)
    ocr_texts = []
    for page in doc:
        for img_info in page.get_images(full=True):
            try:
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                img_data = base_image["image"]
                img = Image.open(io.BytesIO(img_data))
                text = pytesseract.image_to_string(img, lang=TESSERACT_LANGS)
                if text.strip():
                    ocr_texts.append(text)
            except Exception as e:
                print(f"Error extracting image from PDF: {e}")
    return "\n".join(ocr_texts)

def extract_hyperlinks_from_pdf(pdf_path):
    """
    Extract hyperlinks from PDF with PyMuPDF.
    """
    doc = fitz.open(pdf_path)
    links = []
    for page in doc:
        for link in page.get_links():
            if "uri" in link:
                links.append(link["uri"])
    return links

def extract_text_from_docx(docx_path):
    """Extract text from DOCX using python-docx."""
    doc = Document(docx_path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)

def extract_hyperlinks_from_docx(docx_path):
    """
    Not all DOCX files store hyperlinks in run.hyperlink.
    This is partial coverage. 
    """
    doc = Document(docx_path)
    links = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.hyperlink:
                links.append(run.hyperlink.target)
    return links

def extract_images_from_docx(docx_path):
    """
    Extract images from DOCX and OCR them.
    """
    doc = Document(docx_path)
    results = []
    for rel in doc.part.rels:
        target = doc.part.rels[rel].target_ref
        if "image" in target:
            try:
                image_data = doc.part.rels[rel].target_part.blob
                img = Image.open(io.BytesIO(image_data))
                text = pytesseract.image_to_string(img, lang=TESSERACT_LANGS)
                if text.strip():
                    results.append(text)
            except Exception as e:
                print(f"Error extracting DOCX image: {e}")
    return "\n".join(results)

def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_text_from_xlsx(xlsx_path):
    """For Excel files. If CSV, you can parse with read_csv or rename."""
    try:
        df = pd.read_excel(xlsx_path)
        return df.to_string(index=False)
    except Exception as e:
        print(f"Error reading Excel file: {xlsx_path}, {e}")
        return ""

def extract_text_from_csv(csv_path):
    try:
        df = pd.read_csv(csv_path)
        return df.to_string(index=False)
    except Exception as e:
        print(f"Error reading CSV: {csv_path}, {e}")
        return ""

def extract_text_from_image(image_path):
    """OCR for standalone images."""
    try:
        img = Image.open(image_path)
        return pytesseract.image_to_string(img, lang=TESSERACT_LANGS)
    except Exception as e:
        print(f"Error OCR on image: {e}")
        return ""

# -------------------------------------------------------------------
# 4) OPTIONAL YOLO + AI REFINEMENT
# -------------------------------------------------------------------

def detect_objects_in_image(image_path):
    objects_found = []
    try:
        results = model(image_path)
        for r in results:
            if r.boxes is not None:
                for box in r.boxes:
                    class_id = int(box.cls[0].item())
                    class_name = results[0].names[class_id]
                    objects_found.append(class_name)
    except Exception as e:
        print(f"Error running YOLO: {e}")
    return objects_found

def refine_yolo_output_with_ai(objects_detected):
    if not objects_detected:
        return ""
    prompt = (
        "You are an AI summarizing objects in an image.\n"
        f"Objects detected: {', '.join(objects_detected)}\n"
        "Provide a short descriptive summary.but if it is a flow chart please describe properly in full length"
    )
    try:
        response = openai.ChatCompletion.create(
            model=OPENAI_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5
        )
        return response["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"Error refining YOLO with AI: {e}")
        return ""

# -------------------------------------------------------------------
# 5) PLACEHOLDERS (PPT, Video, Google Docs, etc.)
# -------------------------------------------------------------------

def extract_from_pptx(pptx_path):
    # placeholder
    return "Placeholder for PPTX extraction."

def extract_from_video(video_path):
    # placeholder
    return "Placeholder for video transcript."

def extract_from_google_docs(file_id):
    # placeholder
    return "Placeholder for Google Docs text."

# -------------------------------------------------------------------
# 6) MAIN EXTRACTION LOGIC (NO CHUNKING)
# -------------------------------------------------------------------

def process_file(file_path, file_info):
    """
    Extracts raw text, images text, hyperlinks, detected objects, etc.
    Returns a dictionary with:
    {
      "text": "",
      "images_text": "",
      "hyperlinks": [],
      "detected_objects": [],
      "objects_ai_refined": ""
    }
    """
    mime_type = file_info["mimeType"]
    file_name = file_info["name"].lower()

    extracted_data = {
        "text": "",
        "images_text": "",
        "hyperlinks": [],
        "detected_objects": [],
        "objects_ai_refined": ""
    }

    try:
        if mime_type == "application/pdf":
            extracted_data["text"] = extract_text_from_pdf(file_path)
            extracted_data["images_text"] = extract_images_from_pdf(file_path)
            extracted_data["hyperlinks"] = extract_hyperlinks_from_pdf(file_path)

        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            extracted_data["text"] = extract_text_from_docx(file_path)
            extracted_data["images_text"] = extract_images_from_docx(file_path)
            extracted_data["hyperlinks"] = extract_hyperlinks_from_docx(file_path)

        elif mime_type == "application/msword" or file_name.endswith(".doc"):
            # Typically use 'textract' or other tools for .doc
            # For now, fallback
            extracted_data["text"] = "Placeholder for .doc extraction. Use e.g. textract."
            extracted_data["hyperlinks"] = regex_link_extractor(extracted_data["text"])

        elif mime_type in ["text/csv", "application/vnd.ms-excel"] or file_name.endswith(".csv"):
            extracted_data["text"] = extract_text_from_csv(file_path)
            extracted_data["hyperlinks"] = regex_link_extractor(extracted_data["text"])

        elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            extracted_data["text"] = extract_text_from_xlsx(file_path)
            extracted_data["hyperlinks"] = regex_link_extractor(extracted_data["text"])

        elif mime_type == "text/plain":
            extracted_data["text"] = extract_text_from_txt(file_path)
            extracted_data["hyperlinks"] = regex_link_extractor(extracted_data["text"])

        elif mime_type.startswith("image/"):
            # OCR + YOLO
            extracted_data["text"] = extract_text_from_image(file_path)
            objects_found = detect_objects_in_image(file_path)
            extracted_data["detected_objects"] = objects_found
            extracted_data["objects_ai_refined"] = refine_yolo_output_with_ai(objects_found)
            extracted_data["hyperlinks"] = regex_link_extractor(extracted_data["text"])

        elif mime_type == "application/vnd.openxmlformats-officedocument.presentationml.presentation":
            # PPTX placeholder
            ppt_text = extract_from_pptx(file_path)
            extracted_data["text"] = ppt_text
            extracted_data["hyperlinks"] = regex_link_extractor(ppt_text)

        elif mime_type.startswith("video/"):
            # Video placeholder
            vid_text = extract_from_video(file_path)
            extracted_data["text"] = vid_text
            extracted_data["hyperlinks"] = regex_link_extractor(vid_text)

        elif mime_type == "application/vnd.google-apps.document":
            # Google Docs placeholder
            doc_text = extract_from_google_docs(file_info["id"])
            extracted_data["text"] = doc_text
            extracted_data["hyperlinks"] = regex_link_extractor(doc_text)

        else:
            print(f"[!] Unsupported or unknown file type: {mime_type}")
            extracted_data["text"] = ""
            extracted_data["hyperlinks"] = []

    except Exception as e:
        print(f"Error processing {file_info['name']}: {e}")

    return extracted_data

# -------------------------------------------------------------------
# 7) PERSISTENCE / LOG
# -------------------------------------------------------------------

def load_processed_files():
    if os.path.exists(PROCESSED_FILES_LOG):
        with open(PROCESSED_FILES_LOG, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_processed_files(data):
    with open(PROCESSED_FILES_LOG, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4)

# -------------------------------------------------------------------
# 8) MAIN (EXTRACTION-ONLY) PIPELINE
# -------------------------------------------------------------------

def main():
    os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)
    os.makedirs(EXTRACTION_OUTPUT_FOLDER, exist_ok=True)

    creds = authenticate_google_drive()
    files, service = fetch_drive_files(creds)

    processed_files = load_processed_files()

    for file_info in files:
        file_id = file_info["id"]
        file_name = file_info["name"]
        mime_type = file_info["mimeType"]
        modified_time = file_info.get("modifiedTime", "")
        source_url = file_info.get("webViewLink", "")

        # Skip if unchanged
        if file_id in processed_files:
            prev_mod_time = processed_files[file_id].get("modifiedTime", "")
            if modified_time <= prev_mod_time:
                print(f"Skipping unchanged file: {file_name}")
                continue

        print(f"\nExtracting: {file_name} ({mime_type})")

        # Download
        try:
            local_path = download_file(service, file_id, file_name)
        except Exception as e:
            print(f"Error downloading {file_name}: {e}")
            continue

        # Extract
        extracted_data = process_file(local_path, file_info)

        # Build final record
        # (No chunking; purely extraction)
        output_record = {
            "file_id": file_id,
            "file_name": file_name,
            "mime_type": mime_type,
            "source_url": source_url,
            "extracted_content": extracted_data
        }

        # Save to JSON (one file per document)
        out_path = os.path.join(EXTRACTION_OUTPUT_FOLDER, f"{file_id}.json")
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(output_record, f, indent=4)
            print(f"Saved extraction result to {out_path}")
        except Exception as e:
            print(f"Error writing JSON for {file_name}: {e}")

        # Mark as processed
        processed_files[file_id] = {
            "file_name": file_name,
            "modifiedTime": modified_time
        }
        save_processed_files(processed_files)

        # Cleanup downloaded file
        try:
            os.remove(local_path)
        except Exception as e:
            print(f"Error cleaning up {local_path}: {e}")

if __name__ == "__main__":
    main()
